import os
import docx

class DOCXParser:
    def __init__(self, file_path: str):
        self.file_path = file_path
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
            
        self.doc = docx.Document(file_path)
        self.title = os.path.basename(file_path).replace('.docx', '')
        self.chapters = []

    def parse(self):
        """
        解析 DOCX 文件。
        借鉴 Legado 的目录解析逻辑：
        1. 优先使用 Word 自带的标题样式 (Heading)。
        2. 如果没有标题样式，使用正则表达式匹配常见的中文小说章节名。
        3. 如果正则也匹配不到，则按固定长度（段落数）进行无规则拆分。
        """
        import re
        
        current_chapter_title = "前言"
        current_chapter_content = []
        chapter_index = 0
        
        # 常见的中文小说章节正则表达式 (借鉴 Legado)
        chapter_pattern = re.compile(
            r'^\s*(?:第)?\s*[0-9零一二三四五六七八九十百千万两]+\s*(?:章|节|回|卷|折|篇|幕|集|部分)\s*.*$|^\s*(?:序章|前言|引言|楔子|附录|后记)\s*$',
            re.IGNORECASE
        )
        
        has_any_heading = False
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            style_name = para.style.name.lower() if para.style else ""
            
            # 判断是否为标题
            is_heading = False
            if 'heading' in style_name or '标题' in style_name:
                is_heading = True
            elif len(text) < 40 and chapter_pattern.match(text):
                is_heading = True
                
            if is_heading:
                has_any_heading = True
                # 保存上一章
                if current_chapter_content or current_chapter_title != "前言":
                    self.chapters.append({
                        "index": chapter_index,
                        "title": current_chapter_title,
                        "level": 1,
                        "elements": [{"type": "text", "content": p} for p in current_chapter_content]
                    })
                    chapter_index += 1
                
                # 开始新一章
                current_chapter_title = text
                current_chapter_content = []
            else:
                current_chapter_content.append(text)
                
        # 保存最后一章
        if current_chapter_content or current_chapter_title != "前言":
            self.chapters.append({
                "index": chapter_index,
                "title": current_chapter_title,
                "level": 1,
                "elements": [{"type": "text", "content": p} for p in current_chapter_content]
            })
            
        # 如果没有识别出任何章节（无规则拆分目录，借鉴 Legado 的 maxLengthWithNoToc 逻辑）
        if not has_any_heading:
            self.chapters = []
            valid_lines = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
            
            # 假设每 100 个段落为一个分段
            paragraphs_per_chapter = 100
            for i in range(0, len(valid_lines), paragraphs_per_chapter):
                chunk = valid_lines[i:i + paragraphs_per_chapter]
                self.chapters.append({
                    "index": i // paragraphs_per_chapter,
                    "title": f"第 {(i // paragraphs_per_chapter) + 1} 章",
                    "level": 1,
                    "elements": [{"type": "text", "content": p} for p in chunk]
                })
            
        return self.chapters

    def get_chapter(self, index: int):
        if 0 <= index < len(self.chapters):
            return self.chapters[index]
        return None

# 测试代码
if __name__ == "__main__":
    # parser = DOCXParser("test.docx")
    # chapters = parser.parse()
    # for ch in chapters:
    #     print(f"[{ch['index']}] {ch['title']} - 长度: {len(ch['content'])}")
    pass
